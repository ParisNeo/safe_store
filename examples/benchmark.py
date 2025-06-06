import pipmaster as pm

pm.ensure_packages(["pandas", "matplotlib", "seaborn", "scikit-learn", "docx"])

from safe_store import SafeStore
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.manifold import TSNE
from sklearn.decomposition import PCA
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import os

# --- python-docx imports ---
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Helper function to add DataFrame to DOCX (same as before) ---
def add_df_to_doc(doc, df, title=None, bold_header=True, index_bold=False):
    if title:
        doc.add_heading(title, level=2)
    df_for_docx = df.reset_index()
    table = doc.add_table(rows=df_for_docx.shape[0] + 1, cols=df_for_docx.shape[1])
    table.style = 'Table Grid'
    for j, col_name in enumerate(df_for_docx.columns):
        cell = table.cell(0, j)
        cell.text = str(col_name)
        if bold_header or (index_bold and j == 0 and df.index.name is not None):
             cell.paragraphs[0].runs[0].font.bold = True
    for i in range(df_for_docx.shape[0]):
        for j in range(df_for_docx.shape[1]):
            cell_value = df_for_docx.iloc[i, j]
            cell_text = str(cell_value)
            try:
                val = float(cell_value)
                if abs(val) > 0.0001 or val == 0:
                    cell_text = f"{val:.4f}"
                else:
                    cell_text = f"{val:.2e}"
            except (ValueError, TypeError):
                pass
            run = table.cell(i + 1, j).paragraphs[0].runs
            if run: run[0].text = cell_text
            else: table.cell(i + 1, j).text = cell_text
            if index_bold and j == 0 and df.index.name is not None :
                 table.cell(i + 1, j).paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

# --- 1. Define the EXPANDED Dataset ---
benchmark_texts = [
    {
        "id": "animals_fox",
        "query": "The quick brown fox jumps over the lazy dog.",
        "highly_similar": "A fast, dark-colored fox leaps above a sleepy canine.",
        "related": "Dogs are known for their loyalty, while foxes are cunning.",
        "dissimilar": "The stock market experienced significant volatility today.",
        "question_form": "Did the quick brown fox jump over the lazy dog?",
        "negation": "The quick brown fox did not jump over the lazy dog."
    },
    {
        "id": "city_capital",
        "query": "Paris is the capital of France.",
        "highly_similar": "The French capital city is Paris.",
        "related": "France is a country located in Western Europe.",
        "dissimilar": "My favorite food is pizza.",
        "question_form": "Is Paris the capital of France?",
        "negation": "Paris is not the capital of France."
    },
    {
        "id": "tech_python",
        "query": "Python is a versatile programming language.",
        "highly_similar": "Many developers use Python due to its flexibility.",
        "related": "Data science often involves using Python libraries like Pandas and NumPy.",
        "dissimilar": "Gardening requires patience and a green thumb.",
        "question_form": "Is Python a versatile programming language?",
        "negation": "Python is not a difficult programming language to learn."
    },
    {
        "id": "weather_sun",
        "query": "The weather is sunny and warm today.",
        "highly_similar": "It's a bright and balmy day.",
        "related": "I should wear sunscreen because of the sun.",
        "dissimilar": "He is reading a fascinating book about ancient history.",
        "question_form": "Is the weather sunny and warm today?",
        "negation": "The weather is not cloudy and cold today."
    },
    {
        "id": "science_gravity",
        "query": "Gravity is the force that attracts two bodies with mass towards each other.",
        "highly_similar": "The force of attraction between two objects possessing mass is known as gravity.",
        "related": "Albert Einstein's theory of general relativity provides a modern description of gravity.",
        "dissimilar": "Photosynthesis is the process used by plants to convert light energy into chemical energy.",
        "question_form": "Is gravity the force that attracts two bodies with mass towards each other?",
        "negation": "Gravity is not a repulsive force between masses."
    },
    {
        "id": "history_rome",
        "query": "The Roman Empire was one of the most powerful economic, cultural, and military forces in history.",
        "highly_similar": "Among history's dominant powers in economy, culture, and military, the Roman Empire stands out.",
        "related": "Julius Caesar played a critical role in the demise of the Roman Republic and the rise of the Roman Empire.",
        "dissimilar": "The internet has revolutionized global communication.",
        "question_form": "Was the Roman Empire one of history's most powerful forces?",
        "negation": "The Roman Empire was not a minor entity in ancient history."
    },
    {
        "id": "literature_shakespeare",
        "query": "William Shakespeare is widely regarded as the greatest writer in the English language.",
        "highly_similar": "Many consider William Shakespeare the preeminent author in English literature.",
        "related": "Hamlet is one of Shakespeare's most famous tragedies.",
        "dissimilar": "Quantum computing promises to solve complex problems intractable for classical computers.",
        "question_form": "Is William Shakespeare widely regarded as the greatest English writer?",
        "negation": "Few would argue that William Shakespeare was an insignificant writer."
    },
    {
        "id": "tech_ai",
        "query": "Artificial intelligence aims to create machines that can perform tasks requiring human intelligence.",
        "highly_similar": "The goal of AI is to build machines capable of tasks typically needing human intellect.",
        "related": "Machine learning is a subset of artificial intelligence.",
        "dissimilar": "The Amazon rainforest is the world's largest tropical rainforest.",
        "question_form": "Does artificial intelligence aim to create intelligent machines?",
        "negation": "Artificial intelligence is not concerned with developing unintelligent machines."
    },
    {
        "id": "music_mozart",
        "query": "Wolfgang Amadeus Mozart was a prolific and influential composer of the Classical period.",
        "highly_similar": "Mozart, a highly productive and key composer, defined the Classical era of music.",
        "related": "Mozart composed over 800 works, including symphonies, concertos, and operas.",
        "dissimilar": "The stock market closed higher today after a volatile session.", # Re-using a dissimilar
        "question_form": "Was Mozart a prolific composer of the Classical period?",
        "negation": "Mozart was not an obscure composer with little impact."
    },
    {
        "id": "philosophy_ethics",
        "query": "Ethics explores the concepts of right and wrong conduct.",
        "highly_similar": "The study of moral principles concerning good and bad behavior is known as ethics.",
        "related": "Utilitarianism is a theory in normative ethics.",
        "dissimilar": "The recipe calls for two cups of flour and one cup of sugar.",
        "question_form": "Does ethics explore right and wrong conduct?",
        "negation": "Ethics is not indifferent to the distinction between right and wrong."
    },
    {
        "id": "biology_cells",
        "query": "Cells are the basic structural and functional units of all known living organisms.",
        "highly_similar": "The fundamental building blocks of life, in terms of structure and function, are cells.",
        "related": "Eukaryotic cells contain a nucleus, unlike prokaryotic cells.",
        "dissimilar": "Mount Everest is the highest mountain above sea level.",
        "question_form": "Are cells the basic units of living organisms?",
        "negation": "Cells are not irrelevant components of living organisms."
    },
    {
        "id": "economics_supply_demand",
        "query": "Supply and demand is a fundamental concept in economics that describes market dynamics.",
        "highly_similar": "A core economic principle explaining market behavior is supply and demand.",
        "related": "The equilibrium price is where the quantity supplied equals the quantity demanded.",
        "dissimilar": "The artist used vibrant colors to paint the landscape.",
        "question_form": "Is supply and demand a fundamental concept in economics?",
        "negation": "Supply and demand are not minor factors in market analysis."
    },
    {
        "id": "geography_rivers",
        "query": "Rivers play a crucial role in the water cycle and shaping landscapes.",
        "highly_similar": "An essential part of the hydrological cycle and landform development is played by rivers.",
        "related": "The Nile is the longest river in Africa.",
        "dissimilar": "Learning a new language can be a challenging yet rewarding experience.",
        "question_form": "Do rivers play a crucial role in the water cycle?",
        "negation": "Rivers are not insignificant in ecological systems."
    },
    {
        "id": "psychology_memory",
        "query": "Memory is the faculty of the brain by which data or information is encoded, stored, and retrieved when needed.",
        "highly_similar": "The brain's capacity to encode, store, and recall information is called memory.",
        "related": "Short-term memory has a limited capacity compared to long-term memory.",
        "dissimilar": "The new skyscraper will be the tallest in the city.",
        "question_form": "Is memory the brain's faculty for information processing?",
        "negation": "Memory does not involve the immediate forgetting of all information."
    },
    {
        "id": "space_exploration",
        "query": "Space exploration is the use of astronomy and space technology to explore outer space.",
        "highly_similar": "The investigation of outer space using astronomical tools and space tech is known as space exploration.",
        "related": "The Apollo 11 mission was the first to land humans on the Moon.",
        "dissimilar": "Baking a cake requires careful measurement of ingredients.",
        "question_form": "Is space exploration about investigating outer space?",
        "negation": "Space exploration is not focused on studying deep sea trenches."
    },
    {
        "id": "health_exercise",
        "query": "Regular physical exercise is important for maintaining physical fitness and overall health.",
        "highly_similar": "Consistent physical activity is key to preserving fitness and general well-being.",
        "related": "Cardiovascular exercise benefits heart health.",
        "dissimilar": "The museum's new exhibit features ancient artifacts.",
        "question_form": "Is regular physical exercise important for health?",
        "negation": "A sedentary lifestyle is not recommended for good health."
    },
    {
        "id": "food_cooking",
        "query": "Cooking is the art, science, and craft of using heat to prepare food for consumption.",
        "highly_similar": "The practice of preparing food with heat, blending art and science, is called cooking.",
        "related": "Different cooking methods can significantly alter the texture and flavor of ingredients.",
        "dissimilar": "The political debate covered various national issues.",
        "question_form": "Is cooking the art of preparing food with heat?",
        "negation": "Cooking does not involve eating raw ingredients exclusively."
    },
    {
        "id": "climate_change",
        "query": "Climate change refers to long-term shifts in temperatures and weather patterns.",
        "highly_similar": "Significant, extended alterations in temperature and weather define climate change.",
        "related": "Greenhouse gas emissions are a primary driver of current climate change.",
        "dissimilar": "The novel was a bestseller, captivating readers worldwide.",
        "question_form": "Does climate change refer to long-term shifts in weather?",
        "negation": "Climate change is not a short-term, insignificant weather fluctuation."
    },
    {
        "id": "education_learning",
        "query": "Effective learning strategies enhance knowledge retention and understanding.",
        "highly_similar": "Using good study techniques improves how well information is remembered and comprehended.",
        "related": "Active recall and spaced repetition are examples of effective learning strategies.",
        "dissimilar": "The concert featured a renowned symphony orchestra.",
        "question_form": "Do effective learning strategies enhance knowledge retention?",
        "negation": "Poor learning habits do not typically lead to better understanding."
    },
    {
        "id": "food_pizza",
        "query": "I'm craving a delicious slice of pepperoni pizza.",
        "highly_similar": "I really want to eat some tasty pepperoni pizza.",
        "related": "We should order from that new Italian place tonight.",
        "dissimilar": "The documentary about marine biology was fascinating.",
        "question_form": "Am I craving a delicious slice of pepperoni pizza?",
        "negation": "I'm not in the mood for a bland salad right now."
    },
    {
        "id": "chores_laundry",
        "query": "I need to do the laundry because I'm out of clean clothes.",
        "highly_similar": "It's time for laundry since I have no clean garments left.",
        "related": "I hope the washing machine isn't broken again.",
        "dissimilar": "He's planning a trip to the mountains next month.",
        "question_form": "Do I need to do the laundry because I'm out of clean clothes?",
        "negation": "I don't have to worry about washing clothes; my drawers are full."
    },
    {
        "id": "social_meeting",
        "query": "Let's meet for coffee and catch up tomorrow morning.",
        "highly_similar": "How about we get together for coffee tomorrow to chat?",
        "related": "I haven't seen you in ages, so much to talk about!",
        "dissimilar": "The recipe calls for three eggs and a cup of flour.",
        "question_form": "Should we meet for coffee and catch up tomorrow morning?",
        "negation": "We can't get together for dinner tonight; I'm busy."
    },
    {
        "id": "exercise_running",
        "query": "I went for a refreshing run in the park this evening.",
        "highly_similar": "This evening, I enjoyed an invigorating jog through the park.",
        "related": "Running regularly helps me clear my head and stay fit.",
        "dissimilar": "The new software update has some interesting features.",
        "question_form": "Did I go for a refreshing run in the park this evening?",
        "negation": "I stayed home and watched TV all evening instead of exercising."
    },
    {
        "id": "work_deadline",
        "query": "I have a tight deadline to finish this report by Friday.",
        "highly_similar": "This report needs to be completed by Friday, which is a strict deadline.",
        "related": "I'll need to work extra hours to get it done on time.",
        "dissimilar": "She adopted a cute puppy from the animal shelter.",
        "question_form": "Do I have a tight deadline to finish this report by Friday?",
        "negation": "I have plenty of time to complete this report; the deadline isn't for weeks."
    },
    {
        "id": "hobby_gardening",
        "query": "My tomatoes are finally starting to ripen in the garden.",
        "highly_similar": "The tomatoes in my garden are at last beginning to turn red.",
        "related": "I can't wait to make a fresh tomato salad.",
        "dissimilar": "The concert was so loud my ears are still ringing.",
        "question_form": "Are my tomatoes finally starting to ripen in the garden?",
        "negation": "My garden plants are struggling and show no signs of fruit yet."
    },
    {
        "id": "ai_ethics_bias",
        "query": "Ensuring fairness and avoiding bias in AI algorithms is a critical ethical concern.",
        "highly_similar": "It's ethically vital to make AI algorithms fair and prevent them from being biased.",
        "related": "Researchers are developing new techniques to detect and mitigate bias in machine learning models.",
        "dissimilar": "The price of coffee beans has increased due to poor weather conditions in Brazil.",
        "question_form": "Is ensuring fairness and avoiding bias in AI algorithms a critical ethical concern?",
        "negation": "Ignoring potential biases in AI systems is not an acceptable ethical stance."
    },
    {
        "id": "ai_healthcare_diagnosis",
        "query": "AI is being used to improve diagnostic accuracy in medical imaging.",
        "highly_similar": "Artificial intelligence helps enhance the precision of diagnoses from medical scans.",
        "related": "AI algorithms can analyze X-rays and MRIs to identify potential diseases earlier.",
        "dissimilar": "She's learning to play the guitar in her free time.",
        "question_form": "Is AI being used to improve diagnostic accuracy in medical imaging?",
        "negation": "AI is not being ignored as a tool for improving medical diagnostics."
    },
    {
        "id": "ai_nlp_understanding",
        "query": "Natural Language Processing enables computers to understand and generate human language.",
        "highly_similar": "Through NLP, machines gain the ability to comprehend and produce human speech and text.",
        "related": "Chatbots and translation services are common applications of NLP.",
        "dissimilar": "The best way to bake a cake is to follow the recipe carefully.",
        "question_form": "Does Natural Language Processing enable computers to understand human language?",
        "negation": "Natural Language Processing does not make computers incapable of interpreting human communication."
    },
    {
        "id": "ai_ml_learning",
        "query": "Machine learning algorithms allow systems to learn from data without being explicitly programmed.",
        "highly_similar": "Systems can learn from data automatically using machine learning algorithms, without direct programming.",
        "related": "Supervised, unsupervised, and reinforcement learning are different types of machine learning.",
        "dissimilar": "The history of ancient Egypt is filled with fascinating discoveries.",
        "question_form": "Do machine learning algorithms allow systems to learn from data without explicit programming?",
        "negation": "Machine learning systems do not require every single rule to be hard-coded by a programmer."
    },
    {
        "id": "ai_jobs_automation",
        "query": "The impact of AI on the job market and automation is a subject of ongoing debate.",
        "highly_similar": "There's continuous discussion about how AI will affect employment and automation.",
        "related": "Some predict AI will create new jobs, while others worry about widespread job displacement.",
        "dissimilar": "My favorite type of music is classical jazz from the 1950s.",
        "question_form": "Is the impact of AI on the job market a subject of ongoing debate?",
        "negation": "The effect of AI on employment is not a settled matter with no differing opinions."
    },
    {
        "id": "ai_agi_humanlevel",
        "query": "Artificial General Intelligence (AGI) refers to AI that possesses human-like cognitive abilities across diverse tasks.",
        "highly_similar": "AGI is defined as artificial intelligence exhibiting human-level cognitive skills over a wide range of tasks.",
        "related": "Achieving AGI is a long-term goal for many AI researchers, though its feasibility and timeline are debated.",
        "dissimilar": "The local library is hosting a book sale next weekend.",
        "question_form": "Does Artificial General Intelligence refer to AI with human-like cognitive abilities across diverse tasks?",
        "negation": "Artificial General Intelligence is not about creating AI with very narrow, specialized skills limited to one domain."
    },
        {
        "id": "Ancient Forest Legacy",
        "highly_similar": "Descending, we entered a vibrant forest of hardy, evergreen trees and wildlife characteristic of Canada, a landscape shaped by glaciers that retreated 20,000 years ago and pushed this ecosystem further south.",
        "related": "The movement of glaciers across North America has profoundly altered the distribution of plant and animal species, creating unique ecosystems and influencing biodiversity patterns that persist today.",
        "dissimilar": "The development of quantum computing relies on the principles of superposition and entanglement to perform calculations beyond the capabilities of classical computers.",
        "question_form": "How did glacial activity shape the current distribution of Canadian forest ecosystems?",
        "negation": "The forest we entered was not a barren wasteland, but a thriving ecosystem showcasing Canada's natural heritage.",
        "query": "Soon we dropped into a living forest, where cold-tolerant evergreens and boreal animals still evoke the Canadian heritage of an ecosystem pushed south by glaciers 20,000 years ago."
    },
    {
        "id": "Population Growth Rate",
        "highly_similar": "According to the CIA World Factbook's 2011 estimates, the population increased at an annual rate of 1.284%.",
        "related": "Demographic trends significantly impact resource allocation and infrastructure planning, requiring governments to anticipate future population sizes and needs.",
        "dissimilar": "The development of quantum computing relies on the principles of superposition and entanglement to perform complex calculations.",
        "question_form": "What was the estimated annual population growth rate in 2011, according to the CIA World Factbook?",
        "negation": "The CIA World Factbook's 2011 estimates did not indicate a population growth rate exceeding 1.284%.",
        "query": "Annual population growth rate (2011 est., CIA World Factbook): 1.284%."
    },
    {
        "id": "Neonicotinoid Usage Discrepancies",
        "highly_similar": "Due to concerns about their environmental impact, neonicotinoids have been prohibited in the European Union, but their use persists in the United States and Canada.",
        "related": "The debate surrounding pesticide regulation often involves balancing agricultural needs with environmental protection and public health concerns, leading to varying approaches across different countries.",
        "dissimilar": "The development of quantum computing relies on manipulating the superposition and entanglement of qubits to perform complex calculations far beyond the capabilities of classical computers.",
        "question_form": "Why do the United States and Canada continue to use neonicotinoid pesticides while the EU has banned them?",
        "negation": "It is not the case that neonicotinoid pesticides are universally banned across all regions and countries.",
        "query": "This has led to the recent banning of Neonics in the EU, however the US and Canada are still using this chemical pesticide."
    },
    {
        "id": "Color Distribution Across China",
        "highly_similar": "These colors were not limited to specific provinces but were instead distributed unevenly throughout different areas of China.",
        "related": "The inconsistent distribution of pigments in ancient Chinese artifacts suggests complex trade routes and cultural exchange across vast distances, influencing artistic traditions in various regions.",
        "dissimilar": "The development of quantum computing relies on the principles of superposition and entanglement to perform calculations far beyond the capabilities of classical computers.",
        "question_form": "How were these colors distributed geographically across China, and were they limited to specific provinces?",
        "negation": "These colors were not consistently located within a single province, indicating a broader and more dispersed pattern of their presence across China.",
        "query": "In addition, these colors weren't confined to a province but rather irregularly scattered across various regions over all of China."
    },
    {
        "id": "Patient Support During Recovery",
        "highly_similar": "During a patient's recuperation, a loved one or caregiver is permitted to remain with them.",
        "related": "Hospital policies often outline specific guidelines regarding visitor hours and the number of people allowed at a patient's bedside to ensure a comfortable and healing environment for everyone.",
        "dissimilar": "The algorithm efficiently sorts the data, minimizing processing time and maximizing performance.",
        "question_form": "Is it permissible for a family member or support person to stay with a patient while they are recovering?",
        "negation": "A patient is not required to have a family member or support person stay with them during recovery.",
        "query": "A family member or a support person may stay with a patient during recovery."
    },
    {
        "id": "Nevada Emigrant Supplies",
        "highly_similar": "After a lengthy journey along the trail, emigrants arriving in Nevada found their provisions reduced to the essentials.",
        "related": "The scarcity of resources on the Oregon Trail often led emigrants to rely on bartering and foraging to supplement their dwindling supplies.",
        "dissimilar": "The development of quantum computing relies on manipulating the superposition and entanglement of qubits to perform complex calculations.",
        "question_form": "What was the state of emigrant supplies by the time they reached Nevada after traveling the trail?",
        "negation": "Emigrants arriving in Nevada had not traveled a short distance; their journey had been long enough to significantly deplete their supplies.",
        "query": "By the time emigrants got to what is now Nevada, they had been on the trail long enough that their supplies were down to just the basics."
    },
    {
        "id": "Air Cannon Efficiency",
        "highly_similar": "Improved valve designs enable air cannons to function with reduced energy consumption and are frequently implemented in cement plants' primary crushers.",
        "related": "Hydraulic systems are increasingly being explored as alternatives to pneumatic air cannons in industrial crushing applications, offering potentially higher efficiency and more precise control.",
        "dissimilar": "The intricate dance of a hummingbird's wings, beating up to 80 times per second, is a marvel of natural engineering.",
        "question_form": "How do advancements in valve technology impact the operational efficiency of air cannons used in cement plants?",
        "negation": "It is not the case that air cannons cannot benefit from technological improvements in valve design to reduce energy usage.",
        "query": "Advances in valve technology allow air cannons to operate with lower ... generation, coal, metal and .... installed on a cement plant's primary crusher."
    },
    {
        "id": "Erie Canal Length",
        "highly_similar": "The canal's path extended for 363 miles, connecting the Hudson River on the eastern side of New York State to Lake Erie in the west.",
        "related": "The Erie Canal significantly boosted New York's economy by facilitating the transportation of goods and people between the Atlantic coast and the Great Lakes region.",
        "dissimilar": "The Amazon rainforest is experiencing record levels of deforestation, threatening its biodiversity and contributing to climate change.",
        "question_form": "How long is the Erie Canal, and what geographical features does it connect?",
        "negation": "The canal does not extend beyond Lake Erie or originate further east than the Hudson River.",
        "query": "It stretched 363 miles across New York State from the Hudson River in the east to Lake Erie in the west."
    },
    {
        "id": "Connected Development Spaces",
        "highly_similar": "The project's design will incorporate linked areas, accessible roofs, and various routes for circulation.",
        "related": "Urban planners often prioritize creating pedestrian-friendly environments to encourage community interaction and reduce reliance on vehicles.",
        "dissimilar": "The chef meticulously prepared a complex souffl\u00e9, ensuring the oven temperature remained precisely controlled.",
        "question_form": "What elements will characterize the layout of the new development?",
        "negation": "The development will not feature isolated or inaccessible areas, prioritizing open and interconnected spaces.",
        "query": "The new development will therefore create connected spaces, accessible rooftops, pathways and passages."
    },
    {
        "id": "body_heart_pump",
        "query": "The human heart pumps blood throughout the body, supplying oxygen and nutrients.",
        "highly_similar": "Blood circulation, delivering oxygen and nutrients, is driven by the heart's pumping action.",
        "related": "Regular cardiovascular exercise is important for maintaining a healthy heart.",
        "dissimilar": "The invention of the printing press revolutionized information dissemination.",
        "question_form": "Does the human heart pump blood throughout the body to supply oxygen and nutrients?",
        "negation": "The human heart does not primarily function to digest food."
    },
    {
        "id": "body_lungs_breathing",
        "query": "Lungs are essential organs for respiration, exchanging oxygen and carbon dioxide.",
        "highly_similar": "The primary function of the lungs in breathing is the trade of oxygen for carbon dioxide.",
        "related": "Smoking can severely damage lung tissue and impair breathing.",
        "dissimilar": "The capital city of France is Paris, known for its iconic Eiffel Tower.",
        "question_form": "Are lungs essential organs for respiration, exchanging oxygen and carbon dioxide?",
        "negation": "Lungs are not involved in the filtration of waste products from the blood like kidneys are."
    },
    {
        "id": "body_brain_control",
        "query": "The brain is the control center of the nervous system, responsible for thought, memory, and emotion.",
        "highly_similar": "As the hub of the nervous system, the brain governs thoughts, memories, and feelings.",
        "related": "Neurotransmitters are chemicals that transmit signals between nerve cells in the brain.",
        "dissimilar": "The recipe for a good chocolate cake requires high-quality cocoa powder.",
        "question_form": "Is the brain the control center of the nervous system, responsible for thought, memory, and emotion?",
        "negation": "The brain is not a simple muscle responsible only for physical movement."
    },
    {
        "id": "body_bones_support",
        "query": "The skeletal system, composed of bones, provides structural support and protection for the body.",
        "highly_similar": "Bones form the skeleton, which gives the body its structure and safeguards internal organs.",
        "related": "Calcium and Vitamin D are crucial for strong bone development and maintenance.",
        "dissimilar": "Solar panels convert sunlight into electricity, offering a renewable energy source.",
        "question_form": "Does the skeletal system provide structural support and protection for the body?",
        "negation": "The skeletal system is not primarily responsible for digesting food or circulating blood."
    },
    {
        "id": "body_muscles_movement",
        "query": "Muscles contract and relax to produce movement, maintain posture, and generate heat.",
        "highly_similar": "Movement, posture, and heat generation are achieved through the contraction and relaxation of muscles.",
        "related": "Protein is important for muscle repair and growth after exercise.",
        "dissimilar": "The internet allows for instant communication across vast distances.",
        "question_form": "Do muscles contract and relax to produce movement and maintain posture?",
        "negation": "Muscles are not passive tissues; they actively work to enable bodily functions rather than just storing fat."
    },
    {
        "id": "body_immune_defense",
        "query": "The immune system defends the body against pathogens like bacteria and viruses.",
        "highly_similar": "Protection against disease-causing agents such as bacteria and viruses is the role of the immune system.",
        "related": "Vaccines work by training the immune system to recognize and fight specific pathogens.",
        "dissimilar": "Learning a new language can open up opportunities for travel and cultural understanding.",
        "question_form": "Does the immune system defend the body against pathogens like bacteria and viruses?",
        "negation": "The immune system is not designed to facilitate the absorption of nutrients from food."
    }    
]

# Prepare texts and labels for global dimensionality reduction plots
texts_for_global_plot = []
labels_for_global_plot_set_id = []
labels_for_global_plot_type = []
all_unique_texts_set = set()

for item in benchmark_texts:
    set_id = item["id"]
    item_texts = [
        item["query"], item["highly_similar"], item["related"],
        item["dissimilar"], item["question_form"], item["negation"]
    ]
    item_types = ["query", "highly_similar", "related", "dissimilar", "question_form", "negation"]
    
    texts_for_global_plot.extend(item_texts)
    all_unique_texts_set.update(item_texts)
    labels_for_global_plot_set_id.extend([set_id] * len(item_texts))
    labels_for_global_plot_type.extend(item_types)

all_unique_texts_list = list(all_unique_texts_set)
print(f"Total benchmark sets: {len(benchmark_texts)}")
print(f"Total sentences for global plot: {len(texts_for_global_plot)}")
print(f"Total unique sentences to vectorize: {len(all_unique_texts_list)}")

# --- 2. Helper Function for Cosine Similarity ---
def calculate_cosine_similarity(vec1, vec2):
    return cosine_similarity(vec1.reshape(1, -1), vec2.reshape(1, -1))[0][0]

# --- 3. Benchmarking Logic ---
ss = SafeStore(":memory:")
vectorizers_to_test = [
    "st:all-MiniLM-L6-v2",
    "st:all-MiniLM-L12-v2",
    "st:LaBSE"
]
results_data = []
plot_filenames = {}
metrics_cols = ["sim_highly_similar", "sim_related", "sim_dissimilar", "sim_question", "sim_negation"]
discriminator_cols = ["disc_positive_vs_related", "disc_positive_vs_dissimilar", "disc_related_vs_dissimilar"]

print("Vectorizing texts and calculating similarities...")
model_vector_maps = {}
for model_name in vectorizers_to_test:
    print(f"Pre-vectorizing for model: {model_name} ({len(all_unique_texts_list)} sentences)")
    model_vector_maps[model_name] = {text: ss.vectorize_text(text, model_name) for text in all_unique_texts_list}

for model_name in vectorizers_to_test:
    print(f"\nProcessing Model: {model_name}")
    text_to_vector_map = model_vector_maps[model_name]
    for item in benchmark_texts:
        query_vec = text_to_vector_map[item["query"]]
        sim_hs = calculate_cosine_similarity(query_vec, text_to_vector_map[item["highly_similar"]])
        sim_rel = calculate_cosine_similarity(query_vec, text_to_vector_map[item["related"]])
        sim_dis = calculate_cosine_similarity(query_vec, text_to_vector_map[item["dissimilar"]])
        
        results_data.append({
            "model": model_name,
            "set_id": item["id"],
            "query_text": item["query"],
            "sim_highly_similar": sim_hs,
            "sim_related": sim_rel,
            "sim_dissimilar": sim_dis,
            "sim_question": calculate_cosine_similarity(query_vec, text_to_vector_map[item["question_form"]]),
            "sim_negation": calculate_cosine_similarity(query_vec, text_to_vector_map[item["negation"]]),
            "disc_positive_vs_related": sim_hs - sim_rel,
            "disc_positive_vs_dissimilar": sim_hs - sim_dis,
            "disc_related_vs_dissimilar": sim_rel - sim_dis,
        })

df_results = pd.DataFrame(results_data)
agg_stats = df_results.groupby("model")[metrics_cols + discriminator_cols].agg(['mean', 'sem'])
agg_means = agg_stats.xs('mean', level=1, axis=1)
agg_sems = agg_stats.xs('sem', level=1, axis=1)

print("\n--- Aggregated Benchmark Statistics (Mean and SEM) ---")
# print(agg_stats) # This can be very wide, print means and sems separately
print("\n--- Mean Scores ---")
print(agg_means)
print("\n--- Standard Error of Mean (SEM) ---")
print(agg_sems)


# --- Plotting with Error Bars ---
print("\nGenerating plots...")
plot_filenames = {}

# Bar plot for average similarities with error bars
plot_filenames['avg_sim_err'] = "average_similarities_with_error.png"
fig, ax = plt.subplots(figsize=(14, 8))
means_to_plot_sim = agg_means[metrics_cols]
sems_to_plot_sim = agg_sems[metrics_cols]
means_to_plot_sim.plot(kind='bar', yerr=sems_to_plot_sim, ax=ax, capsize=4, width=0.8)
ax.set_title("Average Similarity Scores by Model and Category (with SEM bars)")
ax.set_ylabel("Mean Cosine Similarity")
ax.set_xlabel("Model")
plt.xticks(rotation=45, ha="right")
plt.legend(title="Similarity Type", bbox_to_anchor=(1.05, 1), loc='upper left')
plt.tight_layout(rect=[0, 0, 0.85, 1]) # Adjust layout to make space for legend
plt.savefig(plot_filenames['avg_sim_err'])
plt.close(fig)

# Bar plot for average discriminator scores with error bars
plot_filenames['avg_disc_err'] = "average_discriminator_scores_with_error.png"
fig, ax = plt.subplots(figsize=(14, 8))
means_to_plot_disc = agg_means[discriminator_cols]
sems_to_plot_disc = agg_sems[discriminator_cols]
means_to_plot_disc.plot(kind='bar', yerr=sems_to_plot_disc, ax=ax, capsize=4, width=0.8)
ax.set_title("Average Discriminator Scores by Model (with SEM bars, Higher is Better)")
ax.set_ylabel("Mean Similarity Difference")
ax.set_xlabel("Model")
plt.xticks(rotation=45, ha="right")
plt.legend(title="Discriminator Metric", bbox_to_anchor=(1.05, 1), loc='upper left')
plt.tight_layout(rect=[0, 0, 0.85, 1]) # Adjust layout
plt.savefig(plot_filenames['avg_disc_err'])
plt.close(fig)


# --- Global PCA and t-SNE Plots (Colored by Subject/set_id) ---
plot_filenames['pca_global'] = {}
plot_filenames['tsne_global'] = {}

unique_set_ids_global = sorted(list(set(labels_for_global_plot_set_id)))
# Using a colormap that handles more distinct colors well
cmap_global = plt.cm.get_cmap('tab20', len(unique_set_ids_global)) # tab20 has 20 distinct colors

print("\nGenerating global PCA and t-SNE plots...")
for model_name in vectorizers_to_test:
    print(f"  For model: {model_name}")
    current_model_vector_map = model_vector_maps[model_name]
    text_embeddings_global = np.array([current_model_vector_map[text] for text in texts_for_global_plot])

    # PCA
    pca_filename = f"pca_global_{model_name.replace('st:', '').replace('/', '_').replace(':', '_')}.png"
    plot_filenames['pca_global'][model_name] = pca_filename
    pca = PCA(n_components=2, random_state=42)
    embeddings_2d_pca = pca.fit_transform(text_embeddings_global[:,0,:])

    fig_pca, ax_pca = plt.subplots(figsize=(12, 10)) # Increased size for more legend items
    for i, set_id_val in enumerate(unique_set_ids_global):
        indices = [k for k, label in enumerate(labels_for_global_plot_set_id) if label == set_id_val]
        ax_pca.scatter(embeddings_2d_pca[indices, 0], embeddings_2d_pca[indices, 1],
                       color=cmap_global(i % cmap_global.N), label=set_id_val, alpha=0.7, s=50) # Use modulo for cmap colors
    ax_pca.set_title(f"PCA of All Benchmark Sentences ({model_name})")
    ax_pca.set_xlabel("PCA Component 1")
    ax_pca.set_ylabel("PCA Component 2")
    ax_pca.legend(title="Subject (Set ID)", bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='small')
    plt.tight_layout(rect=[0, 0, 0.80, 1]) # Adjust layout for legend
    plt.savefig(pca_filename)
    plt.close(fig_pca)

    # t-SNE
    tsne_filename = f"tsne_global_{model_name.replace('st:', '').replace('/', '_').replace(':', '_')}.png"
    plot_filenames['tsne_global'][model_name] = tsne_filename
    
    # Adjust perplexity for the number of samples
    n_samples_global = len(texts_for_global_plot)
    perplexity_val = min(30, n_samples_global - 1) # Common default is 30, ensure it's < n_samples
    if perplexity_val <= 1: # t-SNE perplexity must be > 1
        print(f"Skipping t-SNE for {model_name} due to too few samples ({n_samples_global}) for perplexity {perplexity_val}.")
        continue
    
    print(f"    Running t-SNE with perplexity: {perplexity_val} for {n_samples_global} samples.")
    tsne = TSNE(n_components=2, random_state=42, perplexity=perplexity_val, n_iter=1000, learning_rate='auto', init='pca')
    embeddings_2d_tsne = tsne.fit_transform(text_embeddings_global[:,0,:])

    fig_tsne, ax_tsne = plt.subplots(figsize=(12, 10)) # Increased size
    for i, set_id_val in enumerate(unique_set_ids_global):
        indices = [k for k, label in enumerate(labels_for_global_plot_set_id) if label == set_id_val]
        ax_tsne.scatter(embeddings_2d_tsne[indices, 0], embeddings_2d_tsne[indices, 1],
                        color=cmap_global(i % cmap_global.N), label=set_id_val, alpha=0.7, s=50)
    ax_tsne.set_title(f"t-SNE of All Benchmark Sentences ({model_name})")
    ax_tsne.set_xlabel("t-SNE Component 1")
    ax_tsne.set_ylabel("t-SNE Component 2")
    ax_tsne.legend(title="Subject (Set ID)", bbox_to_anchor=(1.05, 1), loc='upper left', fontsize='small')
    plt.tight_layout(rect=[0, 0, 0.80, 1]) # Adjust layout for legend
    plt.savefig(tsne_filename)
    plt.close(fig_tsne)


# Detailed plots for each text set (as before)
plot_filenames['detailed_sets'] = {}
for set_id_val in df_results['set_id'].unique():
    df_subset = df_results[df_results['set_id'] == set_id_val].set_index('model')
    filename = f"similarities_set_{set_id_val}.png"
    plot_filenames['detailed_sets'][set_id_val] = filename
    
    fig, ax = plt.subplots(figsize=(10, 6))
    df_subset[metrics_cols].plot(kind='bar', ax=ax, width=0.8)
    ax.set_title(f"Similarity Scores for Set: '{set_id_val}'")
    ax.set_ylabel("Cosine Similarity")
    plt.xticks(rotation=45, ha="right")
    plt.legend(title="Similarity Type", bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.tight_layout(rect=[0, 0, 0.85, 1]) # Adjust layout
    plt.savefig(filename)
    plt.close(fig)


# --- 4. Create DOCX Report (Structure remains largely the same, interpretations will be more robust) ---
print("\nGenerating DOCX report...")
doc = Document()
doc.add_heading('Maximized Vectorizer Benchmark Report', 0) # Updated title

# Introduction
doc.add_heading('1. Introduction', level=1)
p_intro = doc.add_paragraph()
p_intro.add_run("This report presents an extensive benchmark comparison of sentence vectorizer models: ")
p_intro.add_run(", ".join(vectorizers_to_test) + ".").italic = True
p_intro.add_run(f" The evaluation uses an expanded dataset of {len(benchmark_texts)} diverse subject sets ({len(texts_for_global_plot)} total sentences) "
                "to robustly assess their ability to discriminate between semantically similar, related, and dissimilar sentences, "
                "and their capacity to group texts by subject matter. "
                "Analyses include statistical measures (mean scores with Standard Error of the Mean - SEM) and "
                "dimensionality reduction visualizations (PCA and t-SNE) of the entire benchmark dataset.")

# Methodology
doc.add_heading('2. Methodology', level=1)
doc.add_paragraph(
    f"A predefined dataset of {len(benchmark_texts)} text quads (query, highly similar, related, dissimilar, question, negation) across diverse subjects was used. "
    "Cosine similarity measured embedding closeness. "
    "Aggregated scores (mean and SEM) are based on performance across these subjects. "
    "PCA and t-SNE were applied to all {len(texts_for_global_plot)} benchmark sentences to visualize the embedding space structure for each model, "
    "with points colored by their original subject (Set ID)."
)

# Aggregated Similarity Scores
doc.add_heading('3. Aggregated Similarity Scores (Mean ± SEM)', level=1)
doc.add_paragraph(
    "The table displays mean similarity scores. SEM (Standard Error of Mean) values, indicative of consistency across subjects, are visualized as error bars in the subsequent plot."
)
add_df_to_doc(doc, agg_means[metrics_cols].sort_values(by='sim_highly_similar', ascending=False), title="Mean Similarity Scores by Model", index_bold=True)
doc.add_paragraph("SEM values for the above means:")
add_df_to_doc(doc, agg_sems[metrics_cols].reindex(agg_means[metrics_cols].sort_values(by='sim_highly_similar', ascending=False).index), title="Standard Error of Mean (SEM) for Similarity Scores", index_bold=True)


doc.add_paragraph("The bar chart below visualizes these mean similarity scores with SEM error bars.")
if os.path.exists(plot_filenames['avg_sim_err']):
    doc.add_picture(plot_filenames['avg_sim_err'], width=Inches(6.5))
else:
    doc.add_paragraph(f"[Image '{plot_filenames['avg_sim_err']}' not found]")
p_interp_avg_sim = doc.add_paragraph()
p_interp_avg_sim.add_run("Interpretation: ").bold = True
p_interp_avg_sim.add_run("Models with taller bars for 'sim_highly_similar' and 'sim_question', and shorter bars for 'sim_dissimilar', "
                         "are generally performing better. Smaller error bars (SEM) suggest more consistent performance across different subjects. "
                         "The relative heights of 'sim_related' and 'sim_negation' provide insights into how models handle these nuances.")

# Aggregated Discriminator Scores
doc.add_heading('4. Aggregated Discriminator Scores (Mean ± SEM)', level=1)
doc.add_paragraph(
    "Discriminator scores quantify the separation between categories. Higher positive values are better. SEM values are visualized in the plot."
)
add_df_to_doc(doc, agg_means[discriminator_cols].sort_values(by='disc_positive_vs_dissimilar', ascending=False), title="Mean Discriminator Scores by Model", index_bold=True)
doc.add_paragraph("SEM values for the above discriminator means:")
add_df_to_doc(doc, agg_sems[discriminator_cols].reindex(agg_means[discriminator_cols].sort_values(by='disc_positive_vs_dissimilar', ascending=False).index), title="Standard Error of Mean (SEM) for Discriminator Scores", index_bold=True)


doc.add_paragraph("The bar chart below visualizes these mean discriminator scores with SEM error bars.")
if os.path.exists(plot_filenames['avg_disc_err']):
    doc.add_picture(plot_filenames['avg_disc_err'], width=Inches(6.5))
else:
    doc.add_paragraph(f"[Image '{plot_filenames['avg_disc_err']}' not found]")
p_interp_avg_disc = doc.add_paragraph()
p_interp_avg_disc.add_run("Interpretation: ").bold = True
p_interp_avg_disc.add_run("Larger positive bars, especially for 'disc_positive_vs_dissimilar', indicate stronger discriminatory power. "
                          "Smaller error bars imply more reliable discrimination across subjects.")

# Global Embedding Space Visualization - PCA
doc.add_heading('5. Global Embedding Space: PCA Visualization', level=1)
doc.add_paragraph(
    "PCA reduces dimensionality while preserving variance. Plots show all benchmark sentences in 2D, colored by subject. "
    "This assesses if models group sentences from the same subject and separate different subjects."
)
for model_name in vectorizers_to_test:
    doc.add_heading(f"PCA for {model_name}", level=2)
    pca_plot_path = plot_filenames['pca_global'].get(model_name)
    if pca_plot_path and os.path.exists(pca_plot_path):
        doc.add_picture(pca_plot_path, width=Inches(6.0)) # Adjust width as needed if legend makes it too wide
    else:
        doc.add_paragraph(f"[PCA Image for {model_name} not found at '{pca_plot_path}']")
    p_interp_pca_model = doc.add_paragraph()
    p_interp_pca_model.add_run(f"Interpretation ({model_name}): ").bold = True
    p_interp_pca_model.add_run("Observe clustering of colors. Well-separated color groups indicate good subject discrimination. "
                               "Overlap suggests the model sees similarity between those subjects. Tighter clusters within a color mean consistent subject representation.")
    doc.add_paragraph()

# Global Embedding Space Visualization - t-SNE
doc.add_heading('6. Global Embedding Space: t-SNE Visualization', level=1)
doc.add_paragraph(
    "t-SNE visualizes local structure and clusters. Plots show all benchmark sentences in 2D, colored by subject."
)
for model_name in vectorizers_to_test:
    doc.add_heading(f"t-SNE for {model_name}", level=2)
    tsne_plot_path = plot_filenames['tsne_global'].get(model_name)
    if tsne_plot_path and os.path.exists(tsne_plot_path):
        doc.add_picture(tsne_plot_path, width=Inches(6.0)) # Adjust as needed
    else:
        doc.add_paragraph(f"[t-SNE Image for {model_name} not found at '{tsne_plot_path}']")
    p_interp_tsne_model = doc.add_paragraph()
    p_interp_tsne_model.add_run(f"Interpretation ({model_name}): ").bold = True
    p_interp_tsne_model.add_run("Look for distinct clusters of same-colored points. Clear separation between color clusters is ideal for subject discrimination. "
                                "t-SNE emphasizes local structure; relative distances between distant clusters are less directly meaningful than the quality of local clustering.")
    doc.add_paragraph()

# Detailed Per-Set Results
doc.add_heading(f'7. Detailed Similarity Scores per Text Set (Examples)', level=1)
doc.add_paragraph(
    "This section provides a sample of raw similarity scores for individual text sets, "
    "allowing for a granular look at performance on specific examples. Due to the large number of sets, only a few are shown here."
)
# Show first few sets as examples
sets_to_show_detail = df_results['set_id'].unique()[:3] # Show details for the first 3 sets
for set_id_val in sets_to_show_detail:
    doc.add_heading(f"Analysis for Set: '{set_id_val}'", level=2)
    query_text_for_set = df_results[df_results['set_id'] == set_id_val]['query_text'].iloc[0]
    doc.add_paragraph(f"Query for this set: \"{query_text_for_set}\"")

    df_subset_doc = df_results[df_results['set_id'] == set_id_val][
        ["model"] + metrics_cols
    ].set_index('model')
    add_df_to_doc(doc, df_subset_doc, title=f"Similarity Scores for '{set_id_val}'", index_bold=True)

    detailed_plot_path = plot_filenames['detailed_sets'].get(set_id_val)
    if detailed_plot_path and os.path.exists(detailed_plot_path):
        doc.add_picture(detailed_plot_path, width=Inches(6.0))
    else:
        doc.add_paragraph(f"[Image for set '{set_id_val}' not found at '{detailed_plot_path}']")
    doc.add_paragraph()
    if set_id_val != sets_to_show_detail[-1]:
        doc.add_page_break()

# Conclusion
doc.add_heading('8. Conclusion', level=1)
p_conclusion = doc.add_paragraph()
top_model_disc_mean = agg_means[discriminator_cols]['disc_positive_vs_dissimilar'].idxmax()
top_score_disc_mean = agg_means[discriminator_cols]['disc_positive_vs_dissimilar'].max()

p_conclusion.add_run(
    f"This extensive benchmark, utilizing {len(benchmark_texts)} diverse subject sets, evaluated vectorizer models on semantic discrimination and subject-based clustering. "
    f"Based on mean 'disc_positive_vs_dissimilar' scores, model ").bold = False
p_conclusion.add_run(f"'{top_model_disc_mean}' ({top_score_disc_mean:.4f})").bold = True
p_conclusion.add_run(" demonstrated the strongest average ability to distinguish highly similar sentences from dissimilar ones. "
                     "The SEM values provided insight into the consistency of this performance across the expanded range of subjects.\n")
p_conclusion.add_run(
    "The global PCA and t-SNE visualizations offered qualitative views of how well each model organizes semantic information across multiple topics. "
    "Models producing clearly separated color clusters in these plots are generally better at distinguishing between different subjects. "
    "The increased dataset size for these visualizations provides a more comprehensive test of this capability.\n"
)
p_conclusion.add_run(
    "The choice of the 'best' model remains task-dependent. This comprehensive analysis, with its increased statistical power due to the larger dataset, "
    "provides a more robust multi-faceted view to aid in model selection."
)

# Save Document
report_filename = "vectorizer_benchmark_report_maximized.docx"
doc.save(report_filename)
print(f"\nDOCX report saved as '{report_filename}'")
print("\nMaximized benchmark and reporting complete.")